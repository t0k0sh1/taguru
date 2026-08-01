"""The DOCX connector (ADR 0007 §7/§8, issue #350): ``.docx`` files into
:class:`~taguru_langchain.ingest_connectors.document.ConnectorDocument`,
preserving paragraphs, heading hierarchy, and tables.

``python-docx`` is kept behind a new optional ``docx`` extra (ADR 0007 §3/§4's
packaging decision, the same one ``pdf.py`` already applies to ``pypdf``) — no
dependency forced on a caller who never ingests DOCX.

The document body is walked with ``python-docx``'s own
``iter_inner_content()`` — the one traversal that yields paragraphs and
top-level tables interleaved in real document order (``document.paragraphs``/
``document.tables`` are each individually ordered, but not against each
other). Headers and footers are never visited: they are separate parts, not
children of ``w:body``, so nothing needs to filter them out.

A heading is a paragraph whose style name matches ``Heading N`` (the
`BabelFish`-translated UI name ``python-docx`` gives every localized
``heading N`` style, case-insensitively, so a non-English Word install still
resolves) — or, when a document uses a custom style with no such name at
all, a paragraph carrying its own ``w:pPr/w:outlineLvl`` (0-based; mapped to
1-based levels the same way Word's own UI numbers them). Each heading is its
own paragraph in ``text``, and its breadcrumb (``"Guide > Installation"``,
since ``sections`` is flat — the same convention ``html.py`` already uses)
becomes a :class:`~.document.SectionEntry`.

A table (top-level or nested) becomes exactly ONE paragraph — its rows
joined with ``\\n``, cells with ``" | "`` — carrying a
``{"kind": "table", "value": "N"}`` locator, ``N`` the table's 1-based
position among its document-order siblings (a nested table gets a dotted
path, ``"3.1"``, under its parent's own number). This is the opposite
trade-off from ``html.py``'s ``fragment`` locator (attached to every ordinary
paragraph so a citation can reach a real in-page anchor): a DOCX has no page
boundary of its own to spend that one-locator-per-paragraph budget on (ADR
0007 §7.2), so this connector spends it on tables instead — an ordinary body
paragraph never carries a locator, which is exactly what makes "this
paragraph has a locator" mean "this paragraph is a table" for a citation
reading this document's `locators`.

No OCR engine, no rasterizer: a document left with no extractable text after
the walk (e.g. an image-only DOCX) degrades to `ocr_required` with empty
text — the same code, and the same "not literally a scan" reuse, `html.py`
already applies to a boilerplate-only page. Encrypted (password-protected)
DOCX files use a different container entirely — MS-OFFCRYPTO wraps the whole
package in an OLE2/CFB structure, not a plain zip — so this connector
recognizes that container's own magic bytes and reports `encrypted` before
ever handing the bytes to ``python-docx``, rather than letting them fail as
`corrupt`. A *restricted-editing* password (``w:documentProtection``, stored
in ``word/settings.xml``) is a different, weaker mechanism that leaves the
package a normal, readable zip; this connector does not — and should not —
treat it as `encrypted`.

Footnote/endnote text, comment text, and text-box content are none of them
reachable through a paragraph's own ``.text`` (``python-docx`` has no read
API for the first three at all, and text-box content lives in a nested
``w:txbxContent`` a paragraph's own run text never descends into) — a
document whose body references any of them is named in a single
`partial_extraction` diagnostic rather than silently short-changed. Nested
``<table>``-in-``<table>`` content the reader cannot express any other way
(a merged cell spanning columns, for instance) is read the way row-major
text naturally reads it: repeated once per spanned position, the same
"no attempt at fully faithful merge-cell layout" trade-off ``html.py``'s own
`_handle_table` already accepts.

``.doc`` (legacy binary) and ``.docm`` (macro-enabled) are both
`unsupported_format` — only the plain ``.docx`` OOXML format is read.
"""

from __future__ import annotations

import hashlib
import io
import re
import zipfile
from pathlib import Path
from typing import Final

from taguru import Locator

try:
    from docx import Document as _DocxDocument
    from docx.document import Document as _DocxDocumentType
    from docx.table import Table as _DocxTable
    from docx.table import _Cell as _DocxCell
    from docx.text.paragraph import Paragraph as _DocxParagraph
except ImportError as _python_docx_import_error:  # pragma: no cover - exercised only
    # without python-docx installed
    _DocxDocument = None  # type: ignore[assignment]
    _DocxDocumentType = None  # type: ignore[assignment, misc]
    _DocxTable = None  # type: ignore[assignment, misc]
    _DocxCell = None  # type: ignore[assignment, misc]
    _DocxParagraph = None  # type: ignore[assignment, misc]
    _PYTHON_DOCX_IMPORT_ERROR: ImportError | None = _python_docx_import_error
else:
    _PYTHON_DOCX_IMPORT_ERROR = None

from .._extract import MAX_PASSAGE_BYTES
from ._structure import byte_len, fit_breadcrumb, sanitize_paragraph_text
from .document import (
    MAX_SECTION_BYTES,
    ConnectorDocument,
    ConnectorMetadata,
    Diagnostic,
    DiagnosticCode,
    FingerprintInputs,
    LocatorEntry,
    SectionEntry,
    options_digest,
)
from .sources import check_source_id, file_source_id

PARSER_NAME: Final = "taguru-docx-connector"
PARSER_VERSION: Final = "1.0.0"

_SUPPORTED_SUFFIXES: Final = frozenset({".docx"})
_CONTENT_TYPE: Final = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# The raw-file cap, checked before any parsing — same two-cap shape
# PdfConnector/HtmlConnector use for the same reason: markup/XML overhead
# means a large-but-sparse document can stay well under the extracted-text
# cap, and a heavily-styled small document's raw bytes can still exceed a
# naive per-file cap.
_DEFAULT_MAX_FILE_BYTES: Final = 64 * 1024 * 1024

# MS-OFFCRYPTO password-protected Office documents (including .docx) are
# re-packaged as a whole OLE2/Compound File Binary container, not a zip —
# this is that container format's own magic number (the first 8 bytes of
# every CFB file, regardless of what it holds). Checked BEFORE any zip/XML
# parsing is attempted, so a password-protected file is reported `encrypted`
# rather than falling through to python-docx's own zip-open failure and
# being misreported as `corrupt`.
_OLE2_MAGIC: Final = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

# `BabelFish.internal2ui` (python-docx) already translates a built-in
# style's internal name ("heading 1") to its UI name ("Heading 1") — this
# matches that UI name, case-insensitively (a defensive margin for a custom
# style authored directly under a name python-docx's own translation table
# doesn't happen to cover). `[1-9]` alone is enough: `w:outlineLvl` (the
# fallback below) tops out at 8 (0-based), i.e. heading level 9.
_HEADING_STYLE_RE: Final = re.compile(r"^heading\s*([1-9])$", re.IGNORECASE)

# Content this connector's own paragraph walk cannot reach at all — named in
# a single `partial_extraction` diagnostic rather than silently omitted.
# Every marker is a literal tag search over `word/document.xml`'s raw bytes
# (never a full second parse): a footnote/endnote/comment REFERENCE lives in
# the body even though its own text lives in a separate part this connector
# never reads, and a text box's content lives in a nested `w:txbxContent` a
# paragraph's own `.text` never descends into.
_PARTIAL_CONTENT_MARKERS: Final = (
    (b"<w:footnoteReference", "footnote"),
    (b"<w:endnoteReference", "endnote"),
    (b"<w:commentReference", "comment"),
    (b"<w:txbxContent", "text box"),
)

_EMPTY_SHA256: Final = hashlib.sha256(b"").hexdigest()


def _heading_level(paragraph: _DocxParagraph) -> int | None:
    """The 1-based heading level `paragraph` represents, or `None` if it
    isn't a heading — style name first, `w:outlineLvl` fallback, `None` on
    any lookup failure (a malformed/missing style reference degrades to
    "not a heading", never raises; `python-docx` itself already falls back
    to its own "Normal" style for a `w:pStyle` referencing an undefined
    style, so this rarely needs to matter)."""
    try:
        style = paragraph.style
        style_name = style.name if style is not None else None
    except Exception:  # noqa: BLE001 - a malformed style part degrades to
        # "not a heading" like every other lookup here, never raised.
        style_name = None
    if style_name:
        match = _HEADING_STYLE_RE.match(style_name.strip())
        if match:
            return int(match.group(1))
    try:
        paragraph_properties = paragraph._p.pPr  # noqa: SLF001 - python-docx
        # exposes no public API for `w:outlineLvl`; this is the documented
        # escape hatch every python-docx-based tool uses for the handful of
        # properties the high-level API doesn't wrap.
        outline = paragraph_properties.outlineLvl if paragraph_properties is not None else None
    except Exception:  # noqa: BLE001 - degrades to "not a heading"
        outline = None
    if outline is not None and outline.val is not None:
        return int(outline.val) + 1
    return None


class _BodyResult:
    """The one pass over a document's `iter_inner_content()` builds all
    three at once — mirrors `html.py`'s `_Extractor`, one paragraph-indexed
    pass producing `text` plus paragraph-indexed `locators`/`sections`
    together, so they can never drift out of sync with each other."""

    __slots__ = ("paragraphs", "locators", "sections", "first_heading")

    def __init__(self) -> None:
        self.paragraphs: list[str] = []
        self.locators: list[LocatorEntry] = []
        self.sections: list[SectionEntry] = []
        self.first_heading: str | None = None


def _build_body(
    document: _DocxDocumentType,
    *,
    extract_headings: bool,
    heading_separator: str,
    extract_tables: bool,
) -> _BodyResult:
    result = _BodyResult()
    heading_stack: list[tuple[int, str]] = []

    def commit(text: str) -> int | None:
        text = sanitize_paragraph_text(text)
        if not text:
            return None
        index = len(result.paragraphs)
        result.paragraphs.append(text)
        return index

    def handle_heading(paragraph: _DocxParagraph, level: int) -> None:
        label = paragraph.text.strip()
        index = commit(paragraph.text)
        if index is None:
            return
        if result.first_heading is None:
            result.first_heading = label
        if not extract_headings:
            return
        while heading_stack and heading_stack[-1][0] >= level:
            heading_stack.pop()
        heading_stack.append((level, label))
        breadcrumb = fit_breadcrumb(
            [crumb for _, crumb in heading_stack],
            separator=heading_separator,
            max_bytes=MAX_SECTION_BYTES,
        )
        if breadcrumb is not None:
            result.sections.append(SectionEntry(paragraph=index, section=breadcrumb))

    def handle_table(table: _DocxTable, path: list[int]) -> None:
        rows_text = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            if any(cell.strip() for cell in cells):
                rows_text.append(" | ".join(cells))
        if rows_text:
            index = commit("\n".join(rows_text))
            if index is not None and extract_tables:
                result.locators.append(
                    LocatorEntry(
                        paragraph=index,
                        locator=Locator(kind="table", value=".".join(map(str, path))),
                    )
                )

        # Nested tables (a table inside one of this table's own cells) each
        # become their own separate paragraph+locator, dotted under this
        # table's own path — never inlined into this table's one paragraph
        # of text (§7.3: a table's position must stay distinguishable from
        # its container's). A horizontally-merged cell is repeated verbatim
        # across every column position it spans (the same simplification
        # `_handle_table` above accepts for row text); deduplicated here by
        # underlying element identity so a merged cell's own nested table
        # isn't walked — and counted — more than once.
        seen_cells: list[_DocxCell] = []
        for row in table.rows:
            for cell in row.cells:
                if not any(cell._tc is seen._tc for seen in seen_cells):  # noqa: SLF001
                    seen_cells.append(cell)
        nested_index = 0
        for cell in seen_cells:
            for item in cell.iter_inner_content():
                if isinstance(item, _DocxTable):
                    nested_index += 1
                    handle_table(item, [*path, nested_index])

    table_counter = 0
    for item in document.iter_inner_content():
        if isinstance(item, _DocxParagraph):
            level = _heading_level(item)
            if level is not None:
                handle_heading(item, level)
            else:
                commit(item.text)
        elif isinstance(item, _DocxTable):
            table_counter += 1
            handle_table(item, [table_counter])

    return result


def _partial_extraction_message(raw: bytes) -> str | None:
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            body_xml = archive.read("word/document.xml")
    except (KeyError, zipfile.BadZipFile):
        return None
    kinds = [name for marker, name in _PARTIAL_CONTENT_MARKERS if marker in body_xml]
    if not kinds:
        return None
    return f"{', '.join(kinds)} content is present but not extracted by this connector"


def _document_title(document: _DocxDocumentType, first_heading: str | None) -> str | None:
    try:
        title = document.core_properties.title
    except Exception:  # noqa: BLE001 - malformed core properties degrades
        # to "no title", never raised — the same posture
        # PdfConnector._document_title takes for a malformed `/Info`.
        title = None
    if isinstance(title, str) and title.strip():
        return title.strip()
    return first_heading


class DocxConnector:
    """Reference connector for ``.docx`` files (issue #350).

    ``extract_headings`` (default ``True``) controls whether a detected
    heading paragraph's breadcrumb becomes a `SectionEntry` — the heading's
    own text is always kept as an ordinary paragraph in `text` either way.
    ``extract_tables`` (default ``True``) controls whether a table paragraph
    carries its `{"kind": "table"}` locator — the table's row/cell text is
    always extracted as its own paragraph either way, same posture.
    ``heading_separator`` (default ``" > "``) joins a breadcrumb's levels.
    ``max_file_bytes`` (default 64 MiB) caps the raw file size checked
    before any parsing.
    """

    def __init__(
        self,
        *,
        extract_headings: bool = True,
        heading_separator: str = " > ",
        extract_tables: bool = True,
        max_file_bytes: int = _DEFAULT_MAX_FILE_BYTES,
    ) -> None:
        if _DocxDocument is None:
            raise ImportError(
                "DocxConnector requires python-docx — install with pip install "
                '"langchain-taguru[docx]"'
            ) from _PYTHON_DOCX_IMPORT_ERROR
        self._extract_headings = extract_headings
        self._heading_separator = heading_separator
        self._extract_tables = extract_tables
        self._max_file_bytes = max_file_bytes

    @property
    def parser(self) -> str:
        return PARSER_NAME

    @property
    def parser_version(self) -> str:
        return PARSER_VERSION

    @property
    def parse_options_digest(self) -> str:
        """The digest :meth:`read` would stamp into
        ``fingerprint_inputs.parse_options_digest`` for THIS instance's
        effective config, computable without reading anything (ADR 0007
        §6.3, added for issue #351's connector-level checkpoint: it gates
        skipping a re-parse before ever calling :meth:`read`, which needs
        the full candidate fingerprint up front)."""
        return self._options_digest()

    def supports(self, reference: str) -> bool:
        return Path(reference).suffix.lower() in _SUPPORTED_SUFFIXES

    def _options_digest(self) -> str:
        return options_digest(
            {
                "extract_headings": self._extract_headings,
                "heading_separator": self._heading_separator,
                "extract_tables": self._extract_tables,
                "max_file_bytes": self._max_file_bytes,
            }
        )

    def _fingerprint(self, raw_content_sha256: str) -> FingerprintInputs:
        return FingerprintInputs(
            raw_content_sha256=raw_content_sha256,
            parser=self.parser,
            parser_version=self.parser_version,
            parse_options_digest=self._options_digest(),
        )

    def _failure(
        self,
        *,
        source: str,
        display_name: str,
        code: DiagnosticCode,
        message: str,
        raw_content_sha256: str,
        content_type: str | None = _CONTENT_TYPE,
    ) -> ConnectorDocument:
        return ConnectorDocument(
            source=source,
            text="",
            metadata=ConnectorMetadata(
                origin_uri=source, display_name=display_name, content_type=content_type
            ),
            fingerprint_inputs=self._fingerprint(raw_content_sha256),
            diagnostics=(Diagnostic(code=code, message=message, source=source),),
        )

    def read(self, reference: str) -> ConnectorDocument:
        # As in TextFileConnector/PdfConnector: `reference` itself is the
        # source id (ADR 0007 §6.1); `Path` is only ever used for
        # filesystem access below, never for identity.
        path = Path(reference)
        source = file_source_id(reference)
        display_name = path.name

        def failure(
            code: DiagnosticCode,
            message: str,
            raw_content_sha256: str = _EMPTY_SHA256,
            *,
            content_type: str | None = _CONTENT_TYPE,
        ) -> ConnectorDocument:
            return self._failure(
                source=source,
                display_name=display_name,
                code=code,
                message=message,
                raw_content_sha256=raw_content_sha256,
                content_type=content_type,
            )

        source_diagnostic = check_source_id(source)
        if source_diagnostic is not None:
            return failure(source_diagnostic.code, source_diagnostic.message)

        if not self.supports(reference):
            # Unlike every other failure below, the extension mismatch here
            # is affirmative evidence the file is NOT a DOCX — content_type
            # stays unclaimed (None), the same posture HtmlConnector's own
            # `_failure` takes for its own unsupported-format case, rather
            # than asserting a MIME type this connector has no basis for.
            return failure(
                "unsupported_format",
                f"unsupported extension {path.suffix!r} (only .docx)",
                content_type=None,
            )

        try:
            size = path.stat().st_size
        except OSError as error:
            return failure("unreadable", str(error))
        if size > self._max_file_bytes:
            return failure(
                "content_too_large",
                f"{size} bytes exceeds the {self._max_file_bytes}-byte file cap",
            )

        try:
            raw = path.read_bytes()
        except OSError as error:
            return failure("unreadable", str(error))
        # The raw bytes, before parsing — ADR 0007 §5's connector-native
        # checkpoint fingerprint, deliberately distinct from any hash of
        # the extracted text (§6.2).
        raw_content_sha256 = hashlib.sha256(raw).hexdigest()

        if raw.startswith(_OLE2_MAGIC):
            return failure(
                "encrypted",
                "the document requires a password this connector does not have",
                raw_content_sha256,
            )

        try:
            document = _DocxDocument(io.BytesIO(raw))
            body = _build_body(
                document,
                extract_headings=self._extract_headings,
                heading_separator=self._heading_separator,
                extract_tables=self._extract_tables,
            )
        except Exception as error:  # noqa: BLE001 - python-docx raises a
            # grab-bag of zipfile/lxml/KeyError/ValueError for a malformed
            # package, none sharing a common base — every such failure is
            # still this document's own structure failing to parse.
            return failure("corrupt", str(error), raw_content_sha256)

        if not body.paragraphs:
            return failure(
                "ocr_required",
                "no extractable text (an image-only document, or a genuinely empty one)",
                raw_content_sha256,
            )

        text = "\n\n".join(body.paragraphs)
        if byte_len(text) > MAX_PASSAGE_BYTES:
            return failure(
                "content_too_large",
                f"{byte_len(text)} bytes of extracted text exceeds the "
                f"{MAX_PASSAGE_BYTES}-byte passage cap",
                raw_content_sha256,
            )

        diagnostics: list[Diagnostic] = []
        partial_message = _partial_extraction_message(raw)
        if partial_message is not None:
            diagnostics.append(
                Diagnostic(code="partial_extraction", message=partial_message, source=source)
            )

        return ConnectorDocument(
            source=source,
            text=text,
            locators=tuple(body.locators),
            sections=tuple(body.sections),
            metadata=ConnectorMetadata(
                origin_uri=source,
                display_name=display_name,
                title=_document_title(document, body.first_heading),
                content_type=_CONTENT_TYPE,
            ),
            fingerprint_inputs=self._fingerprint(raw_content_sha256),
            diagnostics=tuple(diagnostics),
        )
