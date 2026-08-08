"""DocxConnector — the .docx connector (ADR 0007 §7/§8, issue #350)."""

from __future__ import annotations

from pathlib import Path

import pytest
from taguru import Locator

pytest.importorskip("docx")

from taguru_langchain._extract import MAX_PASSAGE_BYTES, split_paragraphs  # noqa: E402
from taguru_langchain.ingest_connectors import DocxConnector  # noqa: E402

from .._docx import (  # noqa: E402
    comment_reference_para,
    corrupt_docx,
    docx_bytes,
    encrypted_docx,
    endnote_reference_para,
    footnote_reference_para,
    heading,
    outline_heading,
    para,
    table,
    table_with_nested_cell,
    table_with_nested_cells,
    textbox_para,
)


def _write(tmp_path: Path, name: str, data: bytes) -> Path:
    path = tmp_path / name
    path.write_bytes(data)
    return path


def test_headings_and_paragraphs_are_kept_and_sections_are_breadcrumbs(
    tmp_path: Path,
) -> None:
    body = (
        heading("Title Heading", 1)
        + para("Body para one.")
        + heading("Sub Heading", 2)
        + para("Body para two.")
    )
    path = _write(tmp_path, "doc.docx", docx_bytes(body))
    document = DocxConnector().read(str(path))

    assert document.diagnostics == ()
    assert document.text == ("Title Heading\n\nBody para one.\n\nSub Heading\n\nBody para two.")
    assert [(entry.paragraph, entry.section) for entry in document.sections] == [
        (0, "Title Heading"),
        (2, "Title Heading > Sub Heading"),
    ]
    # An ordinary heading/body paragraph never carries a locator — only a
    # table does (§7.2's one-locator-per-paragraph budget is spent there).
    assert document.locators == ()


def test_outline_level_fallback_recognizes_a_non_heading_named_style(
    tmp_path: Path,
) -> None:
    body = outline_heading("Localized Heading", 2) + para("Body.")
    path = _write(tmp_path, "doc.docx", docx_bytes(body))
    document = DocxConnector().read(str(path))

    assert document.diagnostics == ()
    assert [(entry.paragraph, entry.section) for entry in document.sections] == [
        (0, "Localized Heading")
    ]


def test_outline_level_9_is_body_text_not_a_heading(tmp_path: Path) -> None:
    """ECMA-376: ``w:outlineLvl`` value 9 explicitly means "no outline
    level". ``outline_heading(text, 10)`` emits ``w:val="9"`` (the helper
    is 1-based)."""
    body = outline_heading("Plain body paragraph", 10) + para("Body.")
    path = _write(tmp_path, "doc.docx", docx_bytes(body))
    document = DocxConnector().read(str(path))

    assert document.diagnostics == ()
    assert document.sections == ()


def test_table_becomes_one_paragraph_with_a_table_locator(tmp_path: Path) -> None:
    body = para("Intro") + table([["A1", "B1"], ["A2", "B2"]]) + para("Outro")
    path = _write(tmp_path, "doc.docx", docx_bytes(body))
    document = DocxConnector().read(str(path))

    assert document.diagnostics == ()
    assert document.text == "Intro\n\nA1 | B1\nA2 | B2\n\nOutro"
    assert [(entry.paragraph, entry.locator) for entry in document.locators] == [
        (1, Locator(kind="table", value="1"))
    ]


def test_nested_table_gets_its_own_paragraph_and_dotted_locator(tmp_path: Path) -> None:
    body = (
        para("Intro")
        + table_with_nested_cell(
            [["A1", "B1"], ["A2", "B2"]], at=(1, 0), nested_rows=[["N1", "N2"]]
        )
        + para("Outro")
    )
    path = _write(tmp_path, "doc.docx", docx_bytes(body))
    document = DocxConnector().read(str(path))

    assert document.diagnostics == ()
    assert [(entry.paragraph, entry.locator) for entry in document.locators] == [
        (1, Locator(kind="table", value="1")),
        (2, Locator(kind="table", value="1.1")),
    ]


def test_nested_tables_in_different_cells_of_the_same_parent_get_distinct_locators(
    tmp_path: Path,
) -> None:
    """Regression for a numbering bug: a nested table's ordinal must count
    across every cell of its parent table, not reset to 1 for each cell —
    otherwise two distinct nested tables can end up sharing one locator
    value, which would let a citation resolve to the wrong table."""
    body = para("Intro") + table_with_nested_cells(
        [["", "B1"], ["A2", ""]],
        nested={(0, 0): [["NA1"]], (1, 1): [["NB1"]]},
    )
    path = _write(tmp_path, "doc.docx", docx_bytes(body))
    document = DocxConnector().read(str(path))

    assert document.diagnostics == ()
    values = [entry.locator.value for entry in document.locators]
    assert len(values) == len(set(values)), f"duplicate locator values: {values}"
    assert values == ["1", "1.1", "1.2"]


def test_multiple_top_level_tables_are_numbered_in_document_order(tmp_path: Path) -> None:
    body = table([["A"]]) + para("Between") + table([["B"]])
    path = _write(tmp_path, "doc.docx", docx_bytes(body))
    document = DocxConnector().read(str(path))

    assert [entry.locator.value for entry in document.locators] == ["1", "2"]


def test_built_text_resplits_into_exactly_the_paragraphs_the_locators_and_sections_index(
    tmp_path: Path,
) -> None:
    """The invariant this connector's correctness rests on, mirroring
    ``test_pdf_connector.py``'s own: re-running the server's own paragraph
    splitter over ``text`` must yield exactly as many paragraphs as there
    are, with locator/section indices still pointing at the right one."""
    body = (
        heading("Title", 1)
        + para("Body one.")
        + table([["A1", "B1"], ["A2", "B2"]])
        + para("Body two.\nsecond line.")
        + heading("Next", 2)
        + para("Body three.")
    )
    path = _write(tmp_path, "doc.docx", docx_bytes(body))
    document = DocxConnector().read(str(path))

    resplit = split_paragraphs(document.text)
    assert "\n\n".join(resplit) == document.text
    for entry in document.locators:
        assert 0 <= entry.paragraph < len(resplit)
    for entry in document.sections:
        assert 0 <= entry.paragraph < len(resplit)


def test_real_python_docx_writer_output_round_trips(tmp_path: Path) -> None:
    """Proves this connector reads real Word output, not only
    ``tests/_docx.py``'s own hand-rolled dialect: the fixture module never
    goes through ``python-docx``'s own writer (so it can place XML
    ``python-docx`` has no create-API for, like a footnote reference), but
    this one test builds a document the "normal" way."""
    from docx import Document

    writer = Document()
    writer.add_heading("Real Word Heading", level=1)
    writer.add_paragraph("Real Word body paragraph.")
    table_obj = writer.add_table(rows=1, cols=2)
    table_obj.cell(0, 0).text = "X1"
    table_obj.cell(0, 1).text = "Y1"
    path = tmp_path / "real.docx"
    writer.save(str(path))

    document = DocxConnector().read(str(path))

    assert document.diagnostics == ()
    assert "Real Word Heading" in document.text
    assert "Real Word body paragraph." in document.text
    assert "X1 | Y1" in document.text
    assert [(e.paragraph, e.section) for e in document.sections] == [(0, "Real Word Heading")]
    assert [e.locator for e in document.locators] == [Locator(kind="table", value="1")]


@pytest.mark.parametrize(
    ("builder", "expected_kind"),
    [
        (lambda: footnote_reference_para("See note."), "footnote"),
        (lambda: endnote_reference_para("See note."), "endnote"),
        (lambda: comment_reference_para("Commented."), "comment"),
        (lambda: textbox_para("Box text"), "text box"),
    ],
)
def test_unreachable_content_is_named_partial_extraction(
    tmp_path: Path, builder: object, expected_kind: str
) -> None:
    body = para("Intro") + builder()  # type: ignore[operator]
    path = _write(tmp_path, "doc.docx", docx_bytes(body))
    document = DocxConnector().read(str(path))

    assert [d.code for d in document.diagnostics] == ["partial_extraction"]
    assert expected_kind in document.diagnostics[0].message


def test_corrupt_docx_variants_are_reported_as_corrupt(tmp_path: Path) -> None:
    for kind in ("not_zip", "missing_part", "malformed_xml"):
        path = _write(tmp_path, f"broken-{kind}.docx", corrupt_docx(kind))
        document = DocxConnector().read(str(path))

        assert document.text == ""
        assert [d.code for d in document.diagnostics] == ["corrupt"], kind


def test_encrypted_docx_is_reported_encrypted_without_being_opened(tmp_path: Path) -> None:
    path = _write(tmp_path, "enc.docx", encrypted_docx())
    document = DocxConnector().read(str(path))

    assert document.text == ""
    assert [d.code for d in document.diagnostics] == ["encrypted"]


def test_empty_document_is_reported_ocr_required(tmp_path: Path) -> None:
    path = _write(tmp_path, "empty.docx", docx_bytes(""))
    document = DocxConnector().read(str(path))

    assert document.text == ""
    assert [d.code for d in document.diagnostics] == ["ocr_required"]


def test_unsupported_extension_is_reported_without_touching_the_filesystem(
    tmp_path: Path,
) -> None:
    path = _write(tmp_path, "doc.doc", b"whatever")
    document = DocxConnector().read(str(path))

    assert document.text == ""
    assert [d.code for d in document.diagnostics] == ["unsupported_format"]
    # The extension mismatch is affirmative evidence the file isn't a DOCX
    # — content_type stays unclaimed rather than asserting a MIME type this
    # connector has no basis for (the same posture HtmlConnector's own
    # `_failure` already takes).
    assert document.metadata.content_type is None


def test_docm_extension_is_also_unsupported(tmp_path: Path) -> None:
    path = _write(tmp_path, "doc.docm", docx_bytes(para("x")))
    document = DocxConnector().read(str(path))

    assert [d.code for d in document.diagnostics] == ["unsupported_format"]
    assert document.metadata.content_type is None


def test_other_failure_codes_still_claim_the_docx_mime_type(tmp_path: Path) -> None:
    """Unlike `unsupported_format`, every other failure below happens AFTER
    the `.docx` extension already matched, so claiming the DOCX MIME type
    is a reasonable inference from a trusted extension, not asserted
    fact — kept unchanged by the `unsupported_format` fix above."""
    document = DocxConnector().read(str(tmp_path / "missing.docx"))

    assert [d.code for d in document.diagnostics] == ["unreadable"]
    assert document.metadata.content_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_oversized_file_is_reported_without_parsing(tmp_path: Path) -> None:
    path = tmp_path / "big.docx"
    with open(path, "wb") as handle:
        handle.seek(1024)
        handle.write(b"\0")
    document = DocxConnector(max_file_bytes=1023).read(str(path))

    assert document.text == ""
    assert [d.code for d in document.diagnostics] == ["content_too_large"]


def test_oversized_extracted_text_is_reported_content_too_large(tmp_path: Path) -> None:
    huge_paragraph = "x" * (MAX_PASSAGE_BYTES + 1)
    path = _write(tmp_path, "huge.docx", docx_bytes(para(huge_paragraph)))
    document = DocxConnector().read(str(path))

    assert document.text == ""
    assert [d.code for d in document.diagnostics] == ["content_too_large"]


def test_oversized_source_id_is_reported_without_reading_the_file() -> None:
    long_reference = ("x" * 1025) + ".docx"
    document = DocxConnector().read(long_reference)

    assert document.text == ""
    assert [d.code for d in document.diagnostics] == ["source_id_too_long"]


def test_missing_file_is_reported_unreadable(tmp_path: Path) -> None:
    document = DocxConnector().read(str(tmp_path / "missing.docx"))

    assert document.text == ""
    assert [d.code for d in document.diagnostics] == ["unreadable"]


def test_supports_matches_only_docx(tmp_path: Path) -> None:
    connector = DocxConnector()
    assert connector.supports("a.docx") is True
    assert connector.supports("a.DOCX") is True
    assert connector.supports("a.doc") is False
    assert connector.supports("a.docm") is False
    assert connector.supports("a.txt") is False


def test_parser_identity_is_stamped_into_the_fingerprint(tmp_path: Path) -> None:
    path = _write(tmp_path, "doc.docx", docx_bytes(para("Body.")))
    connector = DocxConnector()
    document = connector.read(str(path))

    assert document.fingerprint_inputs.parser == connector.parser
    assert document.fingerprint_inputs.parser_version == connector.parser_version


def test_fingerprint_hashes_the_raw_bytes_and_the_effective_options(tmp_path: Path) -> None:
    import hashlib

    data = docx_bytes(para("Body."))
    path = _write(tmp_path, "doc.docx", data)

    default_document = DocxConnector().read(str(path))
    assert (
        default_document.fingerprint_inputs.raw_content_sha256 == hashlib.sha256(data).hexdigest()
    )

    other_document = DocxConnector(heading_separator=" / ").read(str(path))
    assert (
        other_document.fingerprint_inputs.parse_options_digest
        != default_document.fingerprint_inputs.parse_options_digest
    )
    # Changing an option never touches the raw-bytes hash — the two
    # fingerprint fields answer independent questions (ADR 0007 §5/§6.2).
    assert (
        other_document.fingerprint_inputs.raw_content_sha256
        == default_document.fingerprint_inputs.raw_content_sha256
    )


def test_extract_headings_false_keeps_paragraph_text_but_drops_sections(
    tmp_path: Path,
) -> None:
    body = heading("Title", 1) + para("Body.")
    path = _write(tmp_path, "doc.docx", docx_bytes(body))

    default_document = DocxConnector().read(str(path))
    document = DocxConnector(extract_headings=False).read(str(path))

    assert document.sections == ()
    assert document.text == default_document.text
    assert (
        document.fingerprint_inputs.parse_options_digest
        != default_document.fingerprint_inputs.parse_options_digest
    )


def test_extract_tables_false_keeps_table_text_but_drops_the_locator(
    tmp_path: Path,
) -> None:
    body = para("Intro") + table([["A1", "B1"]])
    path = _write(tmp_path, "doc.docx", docx_bytes(body))

    default_document = DocxConnector().read(str(path))
    document = DocxConnector(extract_tables=False).read(str(path))

    assert document.locators == ()
    assert document.text == default_document.text
    assert (
        document.fingerprint_inputs.parse_options_digest
        != default_document.fingerprint_inputs.parse_options_digest
    )


def test_metadata_title_prefers_core_properties_over_first_heading(
    tmp_path: Path,
) -> None:
    body = heading("Heading Title", 1) + para("Body.")
    path = _write(tmp_path, "doc.docx", docx_bytes(body, title="Explicit Title"))
    document = DocxConnector().read(str(path))

    assert document.metadata.title == "Explicit Title"


def test_metadata_title_falls_back_to_first_heading(tmp_path: Path) -> None:
    body = heading("Heading Title", 1) + para("Body.")
    path = _write(tmp_path, "doc.docx", docx_bytes(body))
    document = DocxConnector().read(str(path))

    assert document.metadata.title == "Heading Title"


def test_metadata_content_type_is_the_ooxml_wordprocessing_mime_type(
    tmp_path: Path,
) -> None:
    path = _write(tmp_path, "doc.docx", docx_bytes(para("Body.")))
    document = DocxConnector().read(str(path))

    assert document.metadata.content_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_missing_python_docx_raises_a_clear_error_at_construction(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import taguru_langchain.ingest_connectors.docx as docx_module

    monkeypatch.setattr(docx_module, "_DocxDocument", None)
    with pytest.raises(ImportError, match="python-docx"):
        DocxConnector()
